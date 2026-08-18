"""生成代码附录 docx — 从 2024-A 提炼的通用版本。

将指定源文件以语法高亮形式写入 Word 文档，可作为论文附录或代码归档。

用法::

    python scripts/make_appendix.py                          # 独立附录 → outputs/附录.docx
    python scripts/make_appendix.py /path/to/论文.docx       # 追加入论文末尾
    python scripts/make_appendix.py -o outputs/code.docx     # 指定输出路径

配置方法：编辑本文件中的 ``APPENDIX_FILES`` 列表即可。
"""

from __future__ import annotations

import argparse
import ast
import sys
from pathlib import Path

# ---------- 以下库为可选依赖，按需安装 ----------
# pip install python-docx pygments

PROJECT_ROOT = Path(__file__).resolve().parent.parent

# ====================================================================
# 配置区 — 按你的项目修改这里
# ====================================================================

# (章, 节标题, 源文件路径, 节选配置或 None)
# 节选配置: {"funcs": [保留的函数], "class_methods": {类名: [保留的方法]}}
APPENDIX_FILES: list[tuple[str, str, str, dict | None]] = [
    # 示例配置，按项目实际结构修改
    # ("A", "TOPSIS 算法 (src/algorithms/topsis.py)", "src/algorithms/topsis.py", None),
    # ("A", "熵权法 (src/algorithms/entropy_weight.py)", "src/algorithms/entropy_weight.py", None),
    # ("B", "矩阵工具 (src/utils/matrix.py)", "src/utils/matrix.py", None),
    # ("B", "绘图工具 (src/utils/plot.py，节选)",
    #  "src/utils/plot.py", {"funcs": ["setup_cjk_font"], "class_methods": {"Plotter": ["line", "bar"]}}),
]

SECTION_TITLES: dict[str, str] = {
    "A": "附录 A  算法核心代码",
    "B": "附录 B  工具与可视化代码",
}

STRUCTURE_TREE = """项目代码结构：
src/
├── algorithms/    算法实现
├── models/        数学模型基类
├── solve/         求解脚本
├── utils/          矩阵工具、数值方法、绘图
└── io/            数据读写

运行环境：Python 3.10+。"""

# ====================================================================
# 外观配置
# ====================================================================

CODE_FONT = "Consolas"
CJK_FONT = "宋体"
CODE_SIZE = 9  # pt

# VS Code Light+ 配色
_COLOR = {
    "keyword": (0x00, 0x00, 0xFF),
    "string": (0xA3, 0x15, 0x15),
    "comment": (0x00, 0x80, 0x00),
    "number": (0x09, 0x86, 0x58),
    "function": (0x79, 0x5E, 0x26),
    "class": (0x26, 0x7F, 0x99),
    "text": (0x1E, 0x1E, 0x1E),
}


def _token_color(tokentype):
    from pygments.token import Token

    if tokentype in (Token.Comment, Token.Comment.Preproc, Token.Literal.String.Doc):
        return _COLOR["comment"], False, True
    if tokentype in Token.Keyword or tokentype in Token.Operator.Word:
        return _COLOR["keyword"], False, False
    if tokentype in Token.Literal.String:
        return _COLOR["string"], False, False
    if tokentype in Token.Literal.Number:
        return _COLOR["number"], False, False
    if tokentype in (Token.Name.Function, Token.Name.Builtin):
        return _COLOR["function"], False, False
    if tokentype in (Token.Name.Class, Token.Name.Decorator):
        return _COLOR["class"], False, False
    if tokentype is Token.Name.Builtin.Pseudo:
        return _COLOR["keyword"], False, False
    return _COLOR["text"], False, False


def _set_fonts(run, size=None, color=None, bold=False, italic=False):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = CODE_FONT
    run.font.size = Pt(size if size is not None else CODE_SIZE)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), CODE_FONT)
    rFonts.set(qn("w:hAnsi"), CODE_FONT)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)


def _shade(paragraph, fill="F2F2F2"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def _add_code(doc, code: str) -> None:
    from pygments import lex
    from pygments.lexers import PythonLexer
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    def new_par():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Pt(0)
        _shade(p)
        return p

    p = new_par()
    for tokentype, text in lex(code, PythonLexer()):
        color, bold, italic = _token_color(tokentype)
        for i, part in enumerate(text.split("\n")):
            if i > 0:
                if not p.runs:
                    _set_fonts(p.add_run(" "))
                p = new_par()
            if part:
                _set_fonts(p.add_run(part), color=color, bold=bold, italic=italic)
    if p is not None and not p.runs:
        _set_fonts(p.add_run(" "))


def _add_heading(doc, text: str, size_pt: float, level: int | None = None,
                 bold: bool = True) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph(style=f"Heading {level}") if level else doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    _set_fonts(run, size=size_pt, bold=bold, color=(0, 0, 0))
    run.font.name = "黑体"
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "黑体")
    rPr.append(rFonts)


def _excerpt(source: str, config: dict) -> str:
    lines = source.splitlines()
    tree = ast.parse(source)
    chunks = []

    body = tree.body
    first_code = next(
        (i for i, node in enumerate(body)
         if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef))),
        len(body),
    )
    for node in body[:first_code]:
        chunks.append("\n".join(lines[node.lineno - 1:node.end_lineno]))

    keep_funcs = set(config.get("funcs", []))
    keep_methods = config.get("class_methods", {})

    for node in body[first_code:]:
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name in keep_funcs:
            chunks.append("\n".join(lines[node.lineno - 1:node.end_lineno]))
        elif isinstance(node, ast.ClassDef) and node.name in keep_methods:
            wanted = set(keep_methods[node.name])
            methods = [m for m in node.body
                       if isinstance(m, (ast.FunctionDef, ast.AsyncFunctionDef)) and m.name in wanted]
            doc = ast.get_docstring(node, clean=False)
            parts = [f"class {node.name}:"]
            if doc:
                parts.append(f'    """{doc}"""')
            for m in methods:
                parts.append("\n".join(lines[m.lineno - 1:m.end_lineno]) + "\n")
            chunks.append("\n".join(parts).rstrip())

    return "\n\n\n".join(chunks) + "\n"


def _build(doc) -> None:
    from docx.enum.text import WD_BREAK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _add_heading(doc, "附  录", 16, level=1)
    _add_heading(doc, STRUCTURE_TREE, 10.5, bold=False)

    seen: set[str] = set()
    counters: dict[str, int] = {}
    for section, title, rel_path, excerpt_cfg in APPENDIX_FILES:
        if section not in seen:
            _add_heading(doc, SECTION_TITLES.get(section, f"附录 {section}"), 14, level=2)
            seen.add(section)
            counters[section] = 0
        counters[section] += 1

        path = PROJECT_ROOT / rel_path
        if not path.exists():
            print(f"警告：找不到 {rel_path}，已跳过", file=sys.stderr)
            continue

        source = path.read_text(encoding="utf-8")
        if excerpt_cfg:
            source = _excerpt(source, excerpt_cfg)

        _add_heading(doc, f"{section}.{counters[section]}  {title}", 12, level=3)
        _add_code(doc, source)
        print(f"已追加：{rel_path}" + ("（节选）" if excerpt_cfg else ""))


def main() -> None:
    ap = argparse.ArgumentParser(description="生成代码附录 docx")
    ap.add_argument("paper", nargs="?", help="成品论文 .docx 路径")
    ap.add_argument("-o", "--output", help="输出路径")
    args = ap.parse_args()

    from docx import Document

    if args.paper:
        paper = Path(args.paper).expanduser().resolve()
        if not paper.exists():
            sys.exit(f"找不到论文文件：{paper}")
        output = (
            Path(args.output).expanduser().resolve()
            if args.output
            else paper.with_name(f"{paper.stem}_含附录{paper.suffix}")
        )
        doc = Document(paper)
        note = "原论文未改动。"
    else:
        output = (
            Path(args.output).expanduser().resolve()
            if args.output
            else PROJECT_ROOT / "outputs" / "附录.docx"
        )
        output.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        note = "可在 Word 中用「插入 → 对象 → 文件中的文字」将其并入论文。"

    _build(doc)
    doc.save(output)
    print(f"\n完成：{output}")
    print(note)


if __name__ == "__main__":
    main()
